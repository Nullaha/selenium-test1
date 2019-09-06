import os
from bs4 import BeautifulSoup
import requests
import time
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver import ActionChains
import selenium.webdriver.support.ui as ui
import re
import docx

driver = webdriver.Chrome('C:\Program Files (x86)\Google\Chrome\Application\chromedriver.exe')
driver.get("http://www.china-data-online.com/CensusMaps/RequestSearch")
driver.maximize_window()  # 浏览器最大化

driver.find_element_by_id("upperLevel1").click()
    # id="su"是百度搜索按钮，click() 是模拟点击
wait01 = ui.WebDriverWait(driver, 20)

driver.find_element_by_xpath("//*[@id='trIndex']/ul/li[3]/span/span[2]").click()
wait02 = ui.WebDriverWait(driver, 20)

#可以循环!!!!!
driver.find_element_by_xpath("//*[@id='trIndex']/ul/li[3]/ul/li[22]/span/span[2]").click()
wait03 = ui.WebDriverWait(driver, 20)

fenye = driver.find_element_by_xpath("//*[@id='trIndex']/ul/li[3]/ul/li[22]/span/span[2]").text
print(fenye)
#!!!!!!!!!!!

num = driver.find_element_by_css_selector('div.search_result_head>div.resultStatWrapper>span:nth-child(3)').text
print(num)
num2 = re.sub("\D", "", num) #??????
all_pages = int(num2) // 20 + 1
print(all_pages)

pages = 0
sites = []
while pages < all_pages:
    a = driver.find_element_by_css_selector("div.search_result_head > #PaginatedNavigation >span").text
    a =re.sub("\D", "", a)
    pages = int(a)

    wait06 = ui.WebDriverWait(driver, 20)
    links = driver.find_elements_by_css_selector(" div.item_content>div.item_name>a")
    for link in links:
        lin = link.get_attribute('href')
        print(lin)
        sites.append(lin)
    time.sleep(2)

    nex = driver.find_element_by_css_selector("div.search_result_head > #PaginatedNavigation > #NextPageLink").click()
    time.sleep(2)
'''
sites_code = []
for i in sites:
    ii = i.split('China/')
    sites_code.append(ii[-1])
print(sites_code)
print('code总数：'+str(len(sites_code)))
'''

'''
#生成 png网址
image_urls = []
urls = []
for i in sites_code:
    url2 = "http://www.china-data-online.com/CensusMaps/RequestExport?exportType=xls&mapID={code}&imagePath=cache/widget/staticMap/91d4fb781d8a5395a4fa7fe3e5a7c961af227ad4f90945b53ce428a378838138ded4152d3d473a4d4e46919a80b3bae4cd3dd1cbb4672e11d396a2f8ef675&mapTitle=".format(code =i)
    url = "http://www.china-data-online.com/CensusMaps/map/China/{code}".format(code = i)
    image_urls.append(url2)
    urls.append(url)
print('图片地址为：')
print(image_urls)
print('提取标题地址为：')
print(urls)
'''

#直接写word里

# 第五步
file = docx.Document()
for word in fenye:
    file.add_paragraph(word)
for word in num:
    file.add_paragraph(word)
for word in sites:
    file.add_paragraph(word)

file.save("D:\Documents\Desktop\公共管理和社会组织22.docx")


